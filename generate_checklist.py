"""Generate Furutec website checklist as Word document.

Organized PAGE BY PAGE → SECTION BY SECTION so the client can walk
through the deliverable in the same order they experience the site.

Status: Done · Partial · Pending · N/A
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/Users/ericcheah/Furutec/Furutec website 3/Furutec_Website_Checklist.docx"

doc = Document()

# ── Page margins ─────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

BRAND = RGBColor(0x2E, 0x31, 0x92)
INK = RGBColor(0x11, 0x11, 0x14)
MUTED = RGBColor(0x6B, 0x6B, 0x73)

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def add_page_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BRAND
    # Coloured underline rule
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after = Pt(10)
    rule = p_rule.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    rule.font.color.rgb = BRAND
    rule.font.size = Pt(6)
    return p

def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = INK
    return p

def add_para(text, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if italic: run.italic = True
    if color: run.font.color.rgb = color
    if size: run.font.size = Pt(size)
    return p

def add_table(rows):
    """rows: list of (item, status, remarks)."""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    widths = (Cm(9.5), Cm(2.4), Cm(5.5))
    for r in table.rows:
        for c, w in zip(r.cells, widths):
            c.width = w

    hdr = table.rows[0].cells
    hdr[0].text = "Section / Element"
    hdr[1].text = "Status"
    hdr[2].text = "Remarks / Notes"
    for cell in hdr:
        set_cell_shading(cell, "2E3192")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

    for i, (item, status, remarks) in enumerate(rows, start=1):
        row = table.rows[i].cells
        row[0].text = item
        row[1].text = status
        row[2].text = remarks
        for c in row:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        status_cell = row[1]
        if status == "Done":
            set_cell_shading(status_cell, "DCEFD6")
        elif status == "Pending":
            set_cell_shading(status_cell, "FFE2C9")
        elif status == "Partial":
            set_cell_shading(status_cell, "FFF4CC")
        elif status == "N/A":
            set_cell_shading(status_cell, "EEEEEE")
        for paragraph in status_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    doc.add_paragraph()

# ── COVER ───────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_run = cover.add_run("FURUTEC WEBSITE")
cover_run.bold = True
cover_run.font.size = Pt(28)
cover_run.font.color.rgb = BRAND

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Content Build Checklist — Page-by-Page")
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = INK

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run(
    "Each row in this checklist corresponds to a section the client "
    "experiences when visiting the website. Items are grouped by page and "
    "in scroll order from top to bottom."
)
meta_run.italic = True
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = MUTED

doc.add_paragraph()
legend = doc.add_paragraph()
legend_run = legend.add_run(
    "Legend  ·  Done = implemented and verified  ·  Partial = partly delivered "
    "(see Remarks)  ·  Pending = awaiting client input or assets  ·  "
    "N/A = removed or not applicable."
)
legend_run.italic = True
legend_run.font.size = Pt(9)
legend_run.font.color.rgb = MUTED

doc.add_paragraph()

# ── PROJECT OVERVIEW ───────────────────────────────────────
add_section_heading("Project Overview")
add_para(
    "Develop a corporate website that is high-technology and user-friendly, "
    "showcasing products, strengthening brand credibility and generating "
    "business inquiries for Furutec Electrical Sdn Bhd."
)
add_para("Target audience: Contractors · Engineers · Developers and industrial clients.",
         italic=True, color=MUTED)

# ── SITEMAP ────────────────────────────────────────────────
add_section_heading("Sitemap")
add_table([
    ("Home (index.html)",                              "Done", ""),
    ("Our Products (product.html)",                    "Done", ""),
    ("Portfolio — Certifications (portfolio.html)",    "Done", ""),
    ("Portfolio — Track Record (portfolio-track-record.html)", "Done", ""),
    ("Get a Quote (quote.html)",                       "Done", ""),
])

# ============================================================
#  HOMEPAGE
# ============================================================
doc.add_page_break()
add_page_title("Page 1 · Homepage  (index.html)")

add_section_heading("Header / Navigation (shared across all pages)")
add_table([
    ("Logo + Busduct System lock-up",                  "Done", ""),
    ("Nav — Home / Our Products / Portfolio",          "Done", ""),
    ("Header phone number — +603-5635 0538",           "Done", ""),
    ("'Get A Quote' CTA button (rounded, sliding-arrow on hover)", "Done", ""),
    ("Mobile hamburger menu",                          "Done", ""),
])

add_section_heading("Hero Section")
add_table([
    ("Background video — cinematic data centre clip, auto-plays on every page load", "Done", "Web-optimised MP4 (11 MB, 1920px H.264 +faststart)."),
    ("Eyebrow tag — 'UNLOCKING BUSDUCT INNOVATION' with orange pulse dot", "Done", ""),
    ("Title — 'Your Path To Efficiency & Reliability'", "Done", ""),
    ("Body copy — industrial / commercial / infrastructure positioning", "Done", ""),
    ("'Learn More' CTA button",                        "Done", ""),
    ("Certification logo bar at the bottom of the hero", "Done", "White rounded card with 'CERTIFIED' label + 7 real cert logos (Intertek, KEMA, DEKRA, UL, ISO, TÜV, SIRIM)."),
])

add_section_heading("Company Profile + Product R&D (2-column section)")
add_table([
    ("Eyebrow — 'COMPANY PROFILE'",                     "Done", ""),
    ("Title — 'Build on expertise. Driven by innovation.'", "Done", ""),
    ("Paragraph — Furutec Electrical, 31 years R&D, certifications", "Done", ""),
    ("'Learn More' button → EITA Resources website",   "Done", "Links to https://www.eita.com.my (parent company)."),
    ("Eyebrow — 'PRODUCT R&D'",                         "Done", ""),
    ("Title — 'Our Product Milestones'",                "Done", ""),
    ("Milestone card 1 — 'Data Centre Application' (Mission Critical tag)", "Done", ""),
    ("Milestone card 2 — 'Indoor Installation'",        "Done", ""),
    ("Milestone card 3 — 'Outdoor Installation'",       "Done", ""),
    ("Real photography for the 3 milestone cards",      "Pending", "Currently external WordPress URLs (furutec.com.my). Replace with locally hosted client photos before launch."),
])

add_section_heading("Engineered in Malaysia (full-screen with intro video)")
add_table([
    ("Section fills the full viewport (min-height 100vh)", "Done", ""),
    ("Furutec logo intro video plays once when section is scrolled into view", "Done", "Muted background; plays full-bleed."),
    ("White overlay fades from 0% → 100% opacity after the video ends", "Done", ""),
    ("Content (title, paragraphs, checklist, button, globe) zooms in from scale(0.6) → scale(1)", "Done", "1.1 s cubic-bezier."),
    ("Magnet scroll-snap — section pulls itself into place when partly visible", "Done", "Engages at ~20% approach from below, 15% past going up."),
    ("Eyebrow — 'ENGINEERED IN MALAYSIA WITH FOUR DECADES OF EXPERTISE'", "Done", ""),
    ("Title — 'Leading Busduct & Power Distribution Manufacturer, Serving Asia Since 1990.'", "Done", ""),
    ("Two-paragraph body copy",                         "Done", ""),
    ("6-item certification checklist (ASTA & KEMA / 250 A-6,300 A / IEC 61439-6 / Fire-Rated / IoT / Site Support)", "Done", ""),
    ("'Schedule A Factory Visit' button",               "Done", ""),
    ("Interactive globe (drag to rotate)",              "Done", "Centred-right, sized at 126% to give the column breathing room."),
])

add_section_heading("Regional Projects (tabbed grid)")
add_table([
    ("Section header / title",                          "Done", ""),
    ("Region tabs with scroll-reveal animation",        "Done", ""),
    ("Project cards rendered on tab click",             "Done", ""),
    ("Final keep / remove decision",                    "Pending", "Awaiting client confirmation. Currently kept on page."),
])

add_section_heading("Our Products (5-card grid)")
add_table([
    ("Eyebrow — 'OUR PRODUCTS'",                        "Done", ""),
    ("Title — 'Complete busduct solutions for every environment.'", "Done", ""),
    ("Description — indoor → data centre distribution copy", "Done", ""),
    ("Card 1 — Indoor Solution",                        "Done", ""),
    ("Card 2 — Outdoor Solution",                       "Done", ""),
    ("Card 3 — Data Centre Solution",                   "Done", ""),
    ("Card 4 — Power Monitoring Solution",              "Done", ""),
    ("Card 5 — Lighting Solution",                     "Done", ""),
    ("Hover state — card turns brand blue, scales 3%, icon and text invert", "Done", ""),
])

add_section_heading("Product Overview — interactive isometric diagram")
add_table([
    ("Eyebrow — 'PRODUCT OVERVIEW'",                    "Done", ""),
    ("Title — 'How A Furutec Busduct System Comes Together'", "Done", ""),
    ("Isometric drawing displayed on the left",         "Done", ""),
    ("5 numbered clickable hotspots overlaying the diagram", "Done", "Hotspot 1 = Outdoor (Cast Resin), 2 = Indoor (Compact Sandwich), 3 = Lighting (LB & PB), 4 = Data Centre (i-DC), 5 = Power Monitoring System."),
    ("Detail card on the right updates when a hotspot is clicked", "Done", "Photo + sub-title + main title + description + 3 bullets + counter (01 / 05)."),
    ("Previous / next pagination buttons inside the detail card", "Done", ""),
])

add_section_heading("Portfolio Teaser (Credential & Track Record CTA)")
add_table([
    ("Dark cinematic background image with gradient overlay", "Done", ""),
    ("Title — 'Credential & Track Record'",             "Done", ""),
    ("Description — 'Where Credential Meets Capability'", "Done", ""),
    ("Single 'Learn More' button → Portfolio page",     "Done", ""),
])

add_section_heading("Inside Furutec — Factory Video")
add_table([
    ("Eyebrow — 'INSIDE FURUTEC'",                       "Done", ""),
    ("Title — 'Get To Know Furutec.'",                   "Done", ""),
    ("Body copy — short company introduction",           "Done", ""),
    ("Video player with click-to-play button + caption", "Done", "Real factory video (assets/Furutec Factory Video.mp4) with native browser controls."),
    ("Overlay fades out on play, fades back in when video ends", "Done", ""),
    ("Magnet scroll-snap applied to this section too",   "Done", ""),
])

add_section_heading("Get In Touch — CTA Form")
add_table([
    ("Eyebrow — 'GET IN TOUCH'",                         "Done", ""),
    ("Title — 'Ready to specify your next busduct project?'", "Done", ""),
    ("Form field — Select Topic (Indoor / Outdoor / Data Centre / Power Monitoring / Others)", "Done", ""),
    ("Form field — Full Name",                           "Done", ""),
    ("Form field — Corporate Email",                     "Done", ""),
    ("Form field — Phone Number",                        "Done", ""),
    ("Form field — Country",                             "Done", ""),
    ("Form field — Project Description",                 "Done", ""),
    ("Form submit button — 'Submit Now'",                "Done", ""),
    ("Form backend — submissions routed to Amy, Vincent & Darren", "Pending", "Backend / SMTP routing to be wired in deployment phase."),
    ("Contact info — Email info@furutec.com.my",         "Done", ""),
    ("Contact info — Phone +603-5635 0538",              "Done", ""),
    ("Contact info — Working Hours 8:30 AM–6:00 PM (Mon–Fri)", "Done", ""),
    ("Map tabs — Subang Jaya HQ + Penang Manufacturing Plant", "Done", ""),
    ("Directions buttons linking to Google Maps",        "Done", ""),
])

# ============================================================
#  PRODUCTS PAGE
# ============================================================
doc.add_page_break()
add_page_title("Page 2 · Our Products  (product.html)")

add_section_heading("Page Hero")
add_table([
    ("Eyebrow — 'OUR PRODUCTS'",                         "Done", ""),
    ("Title — 'Complete busduct solutions for every environment.'", "Done", ""),
    ("Short intro paragraph",                            "Done", ""),
])

add_section_heading("Floating product sub-nav (sticky as user scrolls)")
add_table([
    ("Quick-jump to Indoor / Outdoor / Data Centre / Power Monitoring / Lighting", "Done", ""),
])

add_section_heading("Section 1 — Indoor Solution  (AH & AH-ES Busduct)")
add_table([
    ("Section title — 'INDOOR SOLUTION'",                "Done", ""),
    ("Product image — AH & AH-ES busduct",               "Done", ""),
    ("Eye-catching badge — 'Most Popular' (green gradient + star icon)", "Done", ""),
    ("Title — 'AH & AH-ES BUSDUCT'",                     "Done", ""),
    ("Subtitle — 'Integrated Power Distribution for Every Scale'", "Done", ""),
    ("Description — sandwich-type busduct copy",         "Done", ""),
    ("9 Key Features (incl. '50% Earth Bar or 100% Housing Ground System')", "Done", ""),
    ("Tech Spec card on the right — Current Rating 500 A–6,300 A, IEC 61439-6 & 61439-1, 1,000 V, Cu/Al, Polyester/Epoxy", "Done", "Redesigned card: full-height stretch, gradient background, 4 px brand accent strip, row dividers."),
    ("'Contact Us' + 'Request a Catalog' buttons",       "Done", ""),
    ("Component Diagram Breakdown — title centred above video", "Done", ""),
    ("Component animation video, auto-plays on scroll-into-view", "Done", "assets/products/indoor-diagram video.mp4. Replay button appears when finished."),
])

add_section_heading("Section 2 — Outdoor Solution  (Cast Resin Busduct)")
add_table([
    ("Section title — 'OUTDOOR SOLUTION'",               "Done", ""),
    ("Product image — Cast Resin busduct",               "Done", ""),
    ("Badge — 'Weatherproof' (royal blue gradient + cloud-rain icon)", "Done", ""),
    ("Title — 'CAST RESIN BUSDUCT'",                     "Done", ""),
    ("Subtitle — 'Safe by Design, Powered by Resin'",    "Done", ""),
    ("Description — extreme-environment cast resin copy", "Done", ""),
    ("7 Key Features (BS 6387 / IEC 60331 / JIS C8364 / CNS 14286 / IP68 etc.)", "Done", ""),
    ("Tech Spec — Current 400 A–6,300 A, IP68, IEC 61439-6 & 61439-1, 1,000 V, Cu/Al", "Done", ""),
    ("'Contact Us' + 'Request a Catalog' buttons",       "Done", ""),
    ("Component animation video, auto-plays on scroll-into-view", "Done", "assets/products/outdoor-diagram video.mp4."),
])

add_section_heading("Section 3 — Data Centre Solution  (i-DC Busduct)")
add_table([
    ("Section title — 'DATA CENTRE SOLUTION'",           "Done", ""),
    ("Product image — i-DC busduct",                     "Done", ""),
    ("Badge — 'Featured' (brand orange gradient + sparkle-star icon)", "Done", ""),
    ("Title — 'i-DC BUSDUCT'",                           "Done", ""),
    ("Subtitle — 'High Density Power, Zero Downtime'",   "Done", ""),
    ("Description — AI / Cloud / Edge uptime copy",      "Done", ""),
    ("10 Key Features (IEC 61439-6 → Power Monitoring integration)", "Done", ""),
    ("Tech Spec — Busduct 250/400/630/800 A · TOU 16/32/63/100 A · Cu · Extruded Al-Alloy", "Done", ""),
    ("'Contact Us' + 'Request a Catalog' buttons",       "Done", ""),
    ("Component animation video, auto-plays on scroll-into-view", "Done", "assets/products/datacenter-diagram video.mp4."),
])

add_section_heading("Section 4 — Power Monitoring System")
add_table([
    ("Single centred 820 px column layout (no tech spec per PDF §3.2.4)", "Done", ""),
    ("Section title — 'POWER MONITORING SOLUTION'",      "Done", ""),
    ("Product image — Power Monitoring",                 "Done", ""),
    ("Badge — 'Smart System' (vibrant purple gradient + CPU icon)", "Done", ""),
    ("Title — 'POWER MONITORING SYSTEM'",                "Done", ""),
    ("Subtitle — 'Precision Monitoring, Proven Savings'", "Done", ""),
    ("Description — sensor + HMI + energy / overload prevention copy", "Done", ""),
    ("'Contact Us' + 'Request a Catalog' buttons",       "Done", ""),
])

add_section_heading("Section 5 — Lighting Solution  (LB & PB Busduct)")
add_table([
    ("Section title — 'LIGHTING SOLUTION'",              "Done", ""),
    ("Product image — LB & PB busduct (two views)",      "Done", ""),
    ("Badge — 'Reliability' (teal gradient + shield-check icon)", "Done", ""),
    ("Title — 'LB & PB BUSDUCT'",                        "Done", ""),
    ("Subtitle — 'Where Light Meets Performance'",       "Done", ""),
    ("Description — Lighting Busduct + Power Busduct copy", "Done", ""),
    ("Tech Spec — LB Cu 16–40 A · PB Cu 63–400 A / Al 100–630 A · IP40/54 · IEC 61439-6 & 61439-1 · 50/60 Hz · 230/400 V", "Done", ""),
    ("'Contact Us' + 'Request a Catalog' buttons",       "Done", ""),
])

# ============================================================
#  PORTFOLIO — CERTIFICATIONS
# ============================================================
doc.add_page_break()
add_page_title("Page 3 · Portfolio — Certifications  (portfolio.html)")

add_section_heading("Page Hero")
add_table([
    ("Eyebrow — 'CERTIFICATE CREDENTIAL'",               "Done", ""),
    ("Title — 'Type-tested. Globally certified.'",       "Done", ""),
    ("Description — type-tested + reputable certification bodies copy", "Done", ""),
])

add_section_heading("Certification Logos Grid")
add_table([
    ("Logo 1 — GCC ISO 9001",                            "Done",    "Real logo image."),
    ("Logo 2 — Intertek",                                "Done",    "Real logo image."),
    ("Logo 3 — DEKRA",                                   "Done",    "Real logo image."),
    ("Logo 4 — UL",                                      "Done",    "Real logo image."),
    ("Logo 5 — SIRIM Eco-Label",                         "Done",    "Real logo image."),
    ("Logo 6 — SIRIM Quality Award",                     "Pending", "Styled text placeholder with 'Logo pending' badge until file arrives."),
    ("Logo 7 — Malaysian Brand",                         "Pending", "Styled text placeholder with 'Logo pending' badge until file arrives."),
    ("Logo 8 — KEMA-KEUR",                               "Done",    "Real logo image."),
    ("Logo 9 — TÜV SÜD / PSB Singapore",                 "Pending", "Currently rendering TÜV Rheinland asset. Awaiting client confirmation whether the logo should be TÜV SÜD."),
    ("Logo 10 — Singapore Green Building Product (SGBC)", "Pending", "Styled text placeholder with 'Logo pending' badge until file arrives."),
])

add_section_heading("Closing Block — 'Where Credential Meets Capability'")
add_table([
    ("Closing wayfinding section linking to Track Record", "Pending", "Awaiting client decision: keep as cross-link, or remove."),
])

# ============================================================
#  PORTFOLIO — TRACK RECORD
# ============================================================
doc.add_page_break()
add_page_title("Page 4 · Portfolio — Track Record  (portfolio-track-record.html)")

add_section_heading("Page Hero")
add_table([
    ("Eyebrow — 'TRACK RECORD'",                         "Done", ""),
    ("Title — 'Trusted across every industry.'",         "Done", ""),
    ("Description — power facilities where reliability is not optional copy", "Done", ""),
])

add_section_heading("Industries Grid — 9 cards")
add_table([
    ("Industry 1 — Data Centre · tag 'High Reliability'", "Done", ""),
    ("Industry 2 — Infrastructure (MRT & Train Station) · tag 'Heavy Duty'", "Done", ""),
    ("Industry 3 — Medical & Healthcare · tag 'Mission Critical'", "Done", ""),
    ("Industry 4 — Industrial · tag 'Up to 6,500 A'",    "Done", ""),
    ("Industry 5 — Oil & Gas · tag 'Harsh Environment'", "Done", ""),
    ("Industry 6 — Commercial · tag 'Space Efficient'",  "Done", ""),
    ("Industry 7 — Residential · tag 'Safe & Compact'",  "Done", ""),
    ("Industry 8 — Education · tag 'Energy Efficient'",  "Done", ""),
    ("Industry 9 — Government · tag 'High Compliance'",  "Done", ""),
    ("Real project photography for the 9 industry cards", "Pending", "Currently Unsplash stock. Replace with client-supplied photos before launch."),
])

# ============================================================
#  GET A QUOTE
# ============================================================
doc.add_page_break()
add_page_title("Page 5 · Get a Quote  (quote.html)")

add_section_heading("Page Hero")
add_table([
    ("Eyebrow — 'GET A QUOTE'",                          "Done", ""),
    ("Title — 'Specify Your Next Busduct Project.'",     "Done", ""),
    ("Sub-paragraph — single-line / load / project timeline + 24 hour response copy", "Done", ""),
])

add_section_heading("Contact Form")
add_table([
    ("Form field — Select Topic (Indoor / Outdoor / Data Centre / Power Monitoring / Others)", "Done", ""),
    ("Form field — Full Name",                           "Done", ""),
    ("Form field — Corporate Email",                     "Done", ""),
    ("Form field — Phone Number",                        "Done", ""),
    ("Form field — Country (dropdown)",                  "Done", ""),
    ("Form field — Project Description",                 "Done", ""),
    ("Form submit button — 'Submit Now'",                "Done", ""),
    ("Submissions routed to Amy, Vincent & Darren mailboxes", "Pending", "SMTP / form-handler wiring during deployment."),
])

add_section_heading("Location + Maps")
add_table([
    ("Subang Jaya HQ location with Directions pill",     "Done", ""),
    ("Penang Manufacturing Plant location with Directions pill", "Done", ""),
    ("Map tabs that swap the Google Maps embed",         "Done", ""),
])

add_section_heading("Direct Contact strip (3 cells)")
add_table([
    ("Email — info@furutec.com.my",                      "Done", ""),
    ("Phone — +603-5635 0538",                           "Done", ""),
    ("Working Hours — 8:30 AM–6:00 PM (Mon–Fri)",        "Done", ""),
])

# ============================================================
#  FOOTER  (shared across all 5 pages)
# ============================================================
doc.add_page_break()
add_page_title("Footer  (shared across all pages)")

add_section_heading("Column 1 — Company identity")
add_table([
    ("Logo + 'Furutec Electrical Sdn Bhd' + Company No. 198001003423 (57207-W)", "Done", ""),
    ("Tagline — 'Engineering The Future Of Power Distribution …'", "Done", ""),
    ("Email — info@furutec.com.my",                      "Done", ""),
    ("Phone — +603-5635 0538",                           "Done", ""),
    ("Location — Subang Jaya, Selangor",                 "Done", ""),
])

add_section_heading("Column 2 — Product")
add_table([
    ("Indoor Solution / Outdoor Solution / Data Centre Solution / Power Monitoring Solution", "Done", "4 items per PDF spec."),
])

add_section_heading("Column 3 — Industries")
add_table([
    ("Data Centres / Commercial Buildings / Public Transportation / Medical & Healthcare", "Done", "Links point to Track Record page."),
])

add_section_heading("Column 4 — Support")
add_table([
    ("Technical Docs / Product Catalog / FAQ / Warranty", "Pending", "Items present; awaiting destination URLs from client."),
])

add_section_heading("Column 5 — Company")
add_table([
    ("About Furutec / Key Projects / Certifications / Contact", "Done", ""),
])

add_section_heading("Social media row")
add_table([
    ("Facebook / YouTube / LinkedIn / WhatsApp icons",   "Pending", "Icons present; awaiting client social profile URLs."),
])

add_section_heading("Bottom strip")
add_table([
    ("Privacy Policy · Terms & Conditions · Sitemap",    "Pending", "Placeholder links; need real legal copy + sitemap.xml."),
    ("Copyright — © 2026 Furutec Electrical Sdn Bhd. All Rights Reserved.", "Done", ""),
])

# ============================================================
#  OUTSTANDING ITEMS — CONSOLIDATED SUMMARY
# ============================================================
doc.add_page_break()
add_page_title("Outstanding Items — Awaiting Client Input")
add_para(
    "Everything below is blocked on either a client decision or a client-supplied "
    "asset. Once these are resolved the site is ready for launch.",
    italic=True, color=MUTED,
)

add_table([
    ("Certification logos × 3 — SIRIM Quality Award, Malaysian Brand, Singapore Green Building Product", "Pending", "Showing styled text placeholders with 'Logo pending' badges until logo files arrive. [Cert page, logos 6 / 7 / 10]"),
    ("TÜV brand confirmation — Rheinland (current asset) vs SÜD (suggested by PDF screenshot)", "Pending", "Image renders. Need client sign-off on the correct TÜV brand. [Cert page, logo 9]"),
    ("Header nav additions — About Us / News / Projects / Contact Us", "Pending", "PDF hero screenshot suggests these nav items; main sitemap omits them. Need client decision."),
    ("Regional Projects tabbed grid — keep / redesign / remove", "Pending", "Currently kept on the homepage with scroll-reveal animations."),
    ("Certifications page closing block — 'Where Credential Meets Capability'", "Pending", "Currently present. Decide keep as wayfinding to Track Record, or remove."),
    ("Milestone card images (homepage Product R&D × 3)", "Pending", "Currently external WordPress URLs from furutec.com.my. Move to client-supplied photography hosted locally."),
    ("Industry card images (Track Record × 9)",           "Pending", "Currently Unsplash stock. Replace with client-supplied project photography."),
    ("Contact form backend — route submissions to Amy, Vincent & Darren", "Pending", "SMTP / form-handler wiring during deployment phase."),
    ("Footer Support column URLs — Technical Docs / Product Catalog / FAQ / Warranty", "Pending", "Awaiting destination URLs."),
    ("Footer Social Media URLs — Facebook / YouTube / LinkedIn / WhatsApp", "Pending", "Awaiting client social profile URLs."),
    ("Footer bottom — Privacy Policy / Terms & Conditions / Sitemap", "Pending", "Need real legal copy + sitemap.xml."),
])

# ── PAGE FOOTER — "Built by BETA Social" on every page ────
import os
LOGO_PATH = "/Users/ericcheah/Furutec/Furutec website 3/assets/beta-social-logo.png"
HAS_LOGO = os.path.exists(LOGO_PATH)

for section in doc.sections:
    footer = section.footer
    # Use the first paragraph (added by python-docx) and clear/rewrite it.
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Reset any inherited runs.
    for run in list(p.runs):
        run.text = ""
    txt_run = p.add_run("Built by  ")
    txt_run.font.size = Pt(8)
    txt_run.font.color.rgb = MUTED
    txt_run.italic = True
    if HAS_LOGO:
        # Embed the logo image, height 12 pt (~16 px).
        img_run = p.add_run()
        img_run.add_picture(LOGO_PATH, height=Pt(14))
    else:
        # Fallback: text-only credit until the logo file is supplied.
        fallback = p.add_run("BETA Social")
        fallback.font.size = Pt(8)
        fallback.font.color.rgb = BRAND
        fallback.bold = True

doc.save(OUTPUT)
print(f"Written: {OUTPUT}")
print(f"BETA Social logo embedded in footer: {'Yes' if HAS_LOGO else 'No (using text fallback)'}")
